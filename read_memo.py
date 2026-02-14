import docx
import sys
import os

file_path = r'C:\Users\Jakgrits\project-mgnt\example\MEMO-PM2026-04 ขออนุมัติปรับปรุงจำนวน Mandays Rev2 .docx'

try:
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        sys.exit(1)

    doc = docx.Document(file_path)
    
    print("=== DOCX Content ===")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
            
    print("\n=== Tables ===")
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            print(" | ".join(row_text))
        print("-" * 20)

except Exception as e:
    print(f"Error reading docx: {e}")
